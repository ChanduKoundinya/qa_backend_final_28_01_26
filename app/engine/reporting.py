import os
import io
import re
import pandas as pd
import docx
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg') # Prevent GUI errors on server
import matplotlib.pyplot as plt
import seaborn as sns
from collections import Counter
from datetime import datetime, timezone
import pytz

# --- Helper Function to Add DataFrames to Word ---
def add_df_to_doc(doc, df_input, title=""):
    """
    Helper function to add a Pandas DataFrame to a Word document as a formatted table.
    """
    if isinstance(df_input, pd.Series):
        df = df_input.reset_index()
        if len(df.columns) == 2:
            df.columns = ['Item', 'Count']
        elif len(df.columns) == 1:
            df.columns = ['Value']
    else:
        df = df_input.copy()

    if title:
        doc.add_heading(title, level=3)

    if df.empty:
        doc.add_paragraph("No data to display for this item.")
        return

    # Clean object columns (remove non-printable chars)
    for col_name_table in df.columns:
        if df[col_name_table].dtype == 'object':
            df[col_name_table] = df[col_name_table].astype(str).str.replace(r'[^\x00-\x7F]+', '', regex=True)

    try:
        # Create Table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, col_name_hdr in enumerate(df.columns):
            hdr_cells[i].text = str(col_name_hdr)
            
        # Data Rows
        for index, row_data in df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row_data):
                row_cells[i].text = str(item)
    except Exception as e:
        doc.add_paragraph(f"Error adding table '{title}' to document: {e}")
        doc.add_paragraph("Table data (first 5 rows as string):")
        doc.add_paragraph(df.head().to_string())
    
    doc.add_paragraph() # Add spacing after table

# --- Main Report Generation Function ---
def generate_docx_report(df, output_path,user_tz='UTC'):
    """
    Generates a comprehensive Word report with charts, adhering to the logic
    from the original incident.py file.
    """
    # 1. Configuration & Setup
    TARGET_SCORE = 90.0
    output_dir = os.path.dirname(output_path) # Directory to save temp charts
    
    doc = docx.Document()
    doc.add_heading('Comprehensive Ticket Audit Analysis Report', level=0)
    try:
        # Get exact UTC time right now
        utc_now = datetime.now(timezone.utc)
        # Convert it to the user's specific timezone
        local_tz = pytz.timezone(user_tz)
        local_now = utc_now.astimezone(local_tz)
        # Format it nicely (e.g., "2026-02-17 10:30:45 AM IST")
        formatted_time = local_now.strftime('%Y-%m-%d %I:%M:%S %p %Z')
    except Exception:
        # Fail-safe just in case the timezone string is invalid
        formatted_time = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%SZ')

    # Add the localized time to the top of the Word doc!
    doc.add_paragraph(f"Report generated on: {formatted_time}")
    doc.add_paragraph(f"Target Score for Agent Performance: {TARGET_SCORE}%")
    doc.add_page_break()

    # 2. Data Cleaning & Preprocessing (Applied to the passed DataFrame)
    # We work on a copy to avoid affecting the original DF if it's used elsewhere
    df_audit = df.copy()
    
    # Fill Missing Values
    if 'Agent' in df_audit.columns:
        df_audit['Agent'].fillna('Unknown Agent', inplace=True)
    if 'Training Needs' in df_audit.columns:
        df_audit['Training Needs'].fillna('None Specified', inplace=True)
    if 'Critical Issues' in df_audit.columns:
        df_audit['Critical Issues'].fillna('None Observed', inplace=True)
    if 'Best Practices' in df_audit.columns:
        df_audit['Best Practices'].fillna('None Observed', inplace=True)
    
    # Convert Dates
    if 'Audit Date' in df_audit.columns:
        df_audit['Audit Date'] = pd.to_datetime(df_audit['Audit Date'], errors='coerce')
        # df_audit.dropna(subset=['Audit Date'], inplace=True) # Optional: Decide if you want to drop rows without dates

    # Convert specific columns to string for consistency
    string_cols_audit = [
        'Response SLA Validation', 'Mandatory Fields Check', 'Category/Sub-Category Sync Validation',
        'Resolution Description Validation', 'Work Notes Validation', 'Short Description Validation',
        'Resolution SLA Validation', 'Critical Issues', 'Training Needs', 'Best Practices'
    ]
    for col in string_cols_audit:
        if col in df_audit.columns:
            df_audit[col] = df_audit[col].astype(str)

    # Ensure numeric score
    if 'Overall Score' in df_audit.columns and not pd.api.types.is_numeric_dtype(df_audit['Overall Score']):
        df_audit['Overall Score'] = pd.to_numeric(df_audit['Overall Score'], errors='coerce')

    # --- Section 2.5: Critical Issues & Training Focus ---
    doc.add_heading('Critical Issues & Training Focus', level=1)
    doc.add_paragraph("This section provides a breakdown of critical issues observed during audits and highlights common training needs.")

    # SLA Compliance Text Stats
    if 'Response SLA Validation' in df_audit.columns:
        response_sla_not_met = df_audit[df_audit['Response SLA Validation'] == 'Immediate Retrain'].shape[0]
        doc.add_paragraph(f"{response_sla_not_met} tickets required immediate retraining for Response SLA.")
    
    if 'Resolution SLA Validation' in df_audit.columns:
        resolution_sla_not_met = df_audit[df_audit['Resolution SLA Validation'] == 'Immediate Retrain'].shape[0]
        doc.add_paragraph(f"{resolution_sla_not_met} tickets required immediate retraining for Resolution SLA.")

    # Documentation Issues Stats
    doc.add_heading("Documentation Issues (Work/Resolution Notes)", level=2)
    if 'Resolution Description Validation' in df_audit.columns:
        res_notes_issues = df_audit[df_audit['Resolution Description Validation'].isin(['Needs Training', 'Immediate Retrain'])].shape[0]
        doc.add_paragraph(f"Resolution notes were inadequate or missing for {res_notes_issues} tickets.")
    
    if 'Work Notes Validation' in df_audit.columns:
        work_notes_issues = df_audit[df_audit['Work Notes Validation'].isin(['Needs Training', 'Immediate Retrain'])].shape[0]
        doc.add_paragraph(f"Work notes were inadequate or missing for {work_notes_issues} tickets.")

    # Critical Issues Table
    doc.add_heading("Breakdown of Critical Issues", level=2)
    critical_issues_list = []
    if 'Critical Issues' in df_audit.columns:
        for item in df_audit['Critical Issues']:
            if pd.notna(item) and str(item).lower() not in ['none observed', 'nan', 'none', 'na']:
                # Split by comma or semicolon
                issues = [s.strip() for s in re.split(r',|;', str(item))]
                critical_issues_list.extend(issues)
        
        critical_issue_counts = pd.Series(Counter(critical_issues_list)).sort_values(ascending=False).head(15)
        if not critical_issue_counts.empty:
            add_df_to_doc(doc, critical_issue_counts, "Top Critical Issues")
        else:
            doc.add_paragraph("No specific critical issues were frequently recorded beyond 'None Observed'.")
    else:
        doc.add_paragraph("'Critical Issues' column not found.")

    # Training Needs Table
    doc.add_heading("Breakdown of Training Needs", level=2)
    training_needs_list = []
    if 'Training Needs' in df_audit.columns:
        for item in df_audit['Training Needs']:
            if pd.notna(item) and str(item).lower() not in ['none specified', 'nan', 'none', 'na']:
                needs = [s.strip() for s in re.split(r',|;', str(item))]
                training_needs_list.extend(needs)
        
        training_need_counts = pd.Series(Counter(training_needs_list)).sort_values(ascending=False).head(15)
        if not training_need_counts.empty:
            add_df_to_doc(doc, training_need_counts, "Top Training Needs")
        else:
            doc.add_paragraph("No specific training needs were frequently recorded beyond 'None Specified'.")
    else:
        doc.add_paragraph("'Training Needs' column not found.")

    doc.add_page_break()

    # --- Section 3: Agent Performance Analysis ---
    doc.add_heading('Agent Performance Analysis', level=1)
    
    if 'Agent' in df_audit.columns and 'Overall Score' in df_audit.columns:
        # Calculate Stats
        agent_avg_scores = df_audit.groupby('Agent')['Overall Score'].mean().reset_index()
        agent_avg_scores['Target Met'] = agent_avg_scores['Overall Score'].apply(
            lambda x: 'Met' if pd.notna(x) and x >= TARGET_SCORE else ('Not Met' if pd.notna(x) else 'N/A')
        )
        agent_avg_scores.rename(columns={'Overall Score': 'Avg. Overall Score'}, inplace=True)
        agent_avg_scores = agent_avg_scores.sort_values(by='Avg. Overall Score', ascending=False)
        
        add_df_to_doc(doc, agent_avg_scores, f"Agent Overall Scores vs. {TARGET_SCORE}% Target")

        # Chart: Top 10 Agents
        if not agent_avg_scores.empty:
            doc.add_heading("Top 10 Agents by Average Score", level=2)
            try:
                plt.figure(figsize=(10, 6))
                top_n = agent_avg_scores.nlargest(10, 'Avg. Overall Score')
                sns.barplot(x='Avg. Overall Score', y='Agent', data=top_n, palette='viridis', hue='Agent', dodge=False)
                plt.legend([], [], frameon=False)
                plt.title('Top 10 Agents')
                plt.tight_layout()
                
                chart_path = os.path.join(output_dir, "temp_agent_chart.png")
                plt.savefig(chart_path)
                plt.close()
                
                doc.add_picture(chart_path, width=Inches(6))
                os.remove(chart_path)
            except Exception as e:
                doc.add_paragraph(f"[Chart generation failed: {e}]")
        
        # Tables: Top/Bottom 5
        add_df_to_doc(doc, agent_avg_scores.head(5), "Top 5 Agents")
        add_df_to_doc(doc, agent_avg_scores.tail(5), "Bottom 5 Agents")
    else:
        doc.add_paragraph("Agent or Overall Score data missing.")

    doc.add_page_break()

    # --- Section 4: SLA and Process Compliance ---
    doc.add_heading('SLA and Process Compliance (Overall)', level=1)

    def calculate_compliance(df, col):
        if col not in df.columns or df.empty: return pd.Series(dtype='float64')
        return df[col].value_counts(normalize=True) * 100

    # Compliance Rates
    if 'Response SLA Validation' in df_audit.columns:
        doc.add_heading("Response SLA Compliance (%)", level=2)
        doc.add_paragraph(calculate_compliance(df_audit, 'Response SLA Validation').to_string())

    if 'Resolution SLA Validation' in df_audit.columns:
        doc.add_heading("Resolution SLA Compliance (%)", level=2)
        doc.add_paragraph(calculate_compliance(df_audit, 'Resolution SLA Validation').to_string())

    # Hotspots Chart
    doc.add_heading("Overall Process Adherence Hotspots", level=2)
    validation_cols = [
        'Mandatory Fields Check', 'Category/Sub-Category Sync Validation',
        'Resolution Description Validation', 'Work Notes Validation',
        'Short Description Validation', 'Resolution SLA Validation'
    ]
    
    adherence_issues = {}
    for col in validation_cols:
        if col in df_audit.columns:
            count = df_audit[df_audit[col].isin(['Needs Training', 'Immediate Retrain'])].shape[0]
            adherence_issues[col] = count
            
    if adherence_issues:
        df_adh = pd.DataFrame(list(adherence_issues.items()), columns=['Check Area', 'Issue Count'])
        df_adh = df_adh.sort_values(by='Issue Count', ascending=False)
        
        add_df_to_doc(doc, df_adh, "Process Adherence Issues Count")
        
        if not df_adh.empty:
            try:
                plt.figure(figsize=(10, 6))
                sns.barplot(x='Issue Count', y='Check Area', data=df_adh, palette='mako', hue='Check Area', dodge=False)
                plt.legend([], [], frameon=False)
                plt.title('Process Adherence Hotspots')
                plt.tight_layout()
                
                chart_path = os.path.join(output_dir, "temp_hotspots_chart.png")
                plt.savefig(chart_path)
                plt.close()
                
                doc.add_picture(chart_path, width=Inches(6))
                os.remove(chart_path)
            except Exception as e:
                doc.add_paragraph(f"[Chart generation failed: {e}]")

    doc.add_page_break()

    # --- Section 5: Incident Pattern Analysis ---
    doc.add_heading('Incident Pattern Analysis', level=1)
    
    if 'Category/Sub-Category Sync Validation' in df_audit.columns:
        cat_issues = df_audit[df_audit['Category/Sub-Category Sync Validation'].isin(['Needs Training', 'Immediate Retrain'])]
        
        if not cat_issues.empty:
            doc.add_heading("Category Sync Issues", level=2)
            counts = cat_issues['Category/Sub-Category Sync Validation'].value_counts().reset_index()
            counts.columns = ['Status', 'Count']
            
            add_df_to_doc(doc, counts, "Category Sync Issue Counts")
            
            try:
                plt.figure(figsize=(8, 5))
                sns.barplot(x='Count', y='Status', data=counts, palette='coolwarm', hue='Status', dodge=False)
                plt.legend([], [], frameon=False)
                plt.title('Category Sync Issues')
                plt.tight_layout()
                
                chart_path = os.path.join(output_dir, "temp_cat_chart.png")
                plt.savefig(chart_path)
                plt.close()
                
                doc.add_picture(chart_path, width=Inches(6))
                os.remove(chart_path)
            except Exception as e:
                doc.add_paragraph(f"[Chart generation failed: {e}]")
        else:
            doc.add_paragraph("No significant Category Sync issues found.")

    # Save Final Document
    doc.save(output_path)
    return output_path  